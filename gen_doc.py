"""
Focused demo document: ESP32-P4 terminal UI + system state machine.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── colours ───────────────────────────────────────────────────────
TITLE_BG   = RGBColor(0x00, 0x46, 0x8C)
H1_COLOR   = RGBColor(0x00, 0x46, 0x8C)
H2_COLOR   = RGBColor(0x20, 0x72, 0xB8)
TBL_HDR    = RGBColor(0x00, 0x46, 0x8C)
TBL_ALT    = RGBColor(0xE8, 0xF0, 0xFB)
C_RED      = RGBColor(0xC0, 0x00, 0x00)
C_ORANGE   = RGBColor(0xC5, 0x5A, 0x11)
C_GREEN    = RGBColor(0x10, 0x7C, 0x10)
C_BLUE     = RGBColor(0x00, 0x46, 0x8C)
CODE_BG    = 'F4F4F4'

# ── helpers ───────────────────────────────────────────────────────

def set_cell_bg(cell, color: RGBColor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(color))
    tcPr.append(shd)

def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcB  = OxmlElement('w:tcBorders')
            for side in ('top','left','bottom','right','insideH','insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), 'B8CCE4')
                tcB.append(b)
            tcPr.append(tcB)

def tbl(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, TBL_HDR)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        r.font.size = Pt(10.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            set_cell_bg(cell, TBL_ALT if ri%2==0 else RGBColor(0xFF,0xFF,0xFF))
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_borders(t)
    if widths:
        for row in t.rows:
            for ci, cell in enumerate(row.cells):
                if ci < len(widths):
                    cell.width = Inches(widths[ci])
    return t

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    # filled bar behind heading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(TITLE_BG))
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run('  ' + text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = H2_COLOR
        run.font.size = Pt(12.5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        run.font.size = Pt(10.5)

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), CODE_BG)
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x2D,0x2D,0x2D)

def bullet(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        for r in p.runs:
            r.font.size = Pt(10.5)

def colored_row_note(doc, color: RGBColor, label: str, desc: str):
    """Single-line emphasis row with colored label."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + '  ')
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = color
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)

def sp(doc): doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
doc = Document()
sec = doc.sections[0]
sec.top_margin    = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)

# ── 封面 ──────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('居家老人 AIoT 智能看护系统')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = H1_COLOR

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ESP32-P4 显示终端与系统状态机')
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = H2_COLOR

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('演示讲解文档')
r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x55,0x55,0x55)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 一  ESP32-P4 显示终端界面
# ══════════════════════════════════════════════════════════════════
h1(doc, '一  ESP32-P4 显示终端界面')
sp(doc)

body(doc,
     'ESP32-P4 显示终端是系统的家庭中枢，搭载 7 英寸 1024×600 MIPI-DSI 触摸屏，'
     '运行基于 LVGL 的 HomeCareHub 应用。界面采用三页横向滑动布局，'
     '通过 MQTT over TLS 实时接收小车、摄像头、CSI 传感器数据并驱动界面更新。')

sp(doc)
h2(doc, '1.1  三页界面布局')

tbl(doc,
    ['页面', '核心内容', '实时数据来源'],
    [
        ['Page 1\n概览页',
         '• 顶栏：时间 / 日期 / 看护模式标签 / 温度 / 隐私标签\n'
         '• 左栏：场景描述 + 舒适度仪表盘（温度弧、光照、功率、空气质量）\n'
         '            + CSI 信号波形图（128 点实时折线）\n'
         '• 右栏：天气卡（城市、天气描述、温度、湿度、AQI / CO₂ / 噪声）',
         'MQTT: devices/+/csi/summary\nHTTPS 天气 API（15 分钟刷新）'],
        ['Page 2\n巡逻页',
         '• 左栏：SmartCar 状态面板\n'
         '            — 大字标题显示当前状态\n'
         '            — 5 行元数据（位置/目标/障碍/传感器/语音）\n'
         '            — 3 个地点异常按钮（浴室/卧室/厨房）\n'
         '            — 3 个动作按钮（巡航 / 返航 / 停止）\n'
         '• 右栏：摄像头 JPEG 实时预览（自动缩放适配面板）',
         'MQTT: smartcar/attitude\nMQTT: smartcar/system/status\nMQTT: cam/jpeg'],
        ['Page 3\n控制页',
         '• 左栏：设备在线状态（SmartCar / Camera / CSI 三行）\n'
         '• 右上：自动化日程（3 条固定计划）\n'
         '• 右下：语音助手面板（转写记录、麦克风按钮、隐私切换）',
         'MQTT 在线状态实时更新'],
    ],
    widths=[1.0, 3.7, 2.2]
)
sp(doc)

h2(doc, '1.2  看护模式与界面配色')
body(doc, '终端根据 MQTT 消息中的模式字段实时切换界面主题配色和顶部状态标签：')
sp(doc)

tbl(doc,
    ['模式名称', '触发来源', '主题色', '顶部模式标签', 'SmartCar 状态'],
    [
        ['全屋安全\nNormal',   'cruise / return_running 状态',  '蓝色  #0070F3', '全屋安全',  '待命中 / 返航中'],
        ['异常复核\nFall',     'abnormal_running / ready 状态', '红色  #EE0000', '异常复核',  '前往现场 / 就绪确认'],
        ['浴室看护\nBathroom', '浴室滞留事件触发',              '琥珀  #F5A623', '浴室看护',  '门外等待'],
        ['夜间守护\nNight',    '离床感知事件触发',              '深蓝  #0070F3', '夜间守护',  '低速跟随'],
    ],
    widths=[1.3, 1.8, 1.2, 1.2, 1.4]
)
sp(doc)

h2(doc, '1.3  Page 2 SmartCar 面板详解')
body(doc, 'Page 2 左侧 SmartCar 面板实时响应 MQTT 消息，是演示重点：')
sp(doc)

tbl(doc,
    ['MQTT 消息字段', '面板更新内容'],
    [
        ['smartcar/system/status → state', '大字标题变化：cruise=待命中，abnormal=前往现场，ready=就绪确认，return=返航中'],
        ['smartcar/attitude → r/p/y',      '"SmartCar 在线"标签 + 显示实时 roll / pitch / yaw 数值，副标题改为"LIVE · 姿态回传"'],
        ['smartcar/attitude → battery',    '电池电量百分比显示在面板元数据第 5 行'],
        ['smartcar/system/status → abnormal_ready', '"返航"按钮高亮可用（其他状态下禁用）'],
        ['INBOUND_EVENT → level',          '事件列表顶部插入一条（L3 红色 / L2 橙色 / L1 蓝色）'],
    ],
    widths=[2.4, 5.1]
)
sp(doc)

h2(doc, '1.4  Page 2 摄像头预览面板')
bullet(doc, [
    '显示终端打开 HomeCareHub 应用时，自动调用 camera_mqtt_receiver_set_active(true) 开始接收 JPEG 帧。',
    '后台 JPEG 解码任务将 cam/jpeg 主题的二进制帧解码并渲染到 lv_img 控件，自动缩放适配面板尺寸。',
    '若已连接 Broker 但 5 秒内未收到图传帧，终端自动向 cam/jpeg/control 发布 {"mode":"mqtt"} 触发摄像头切换图传模式。',
    '面板右下角状态标签实时显示分辨率或连接状态。',
])
sp(doc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 二  系统状态机流程
# ══════════════════════════════════════════════════════════════════
h1(doc, '二  系统状态机流程')
sp(doc)

body(doc,
     '系统采用三大主状态（巡航 / 异常 / 返航）组织各模块业务逻辑，'
     '所有状态切换通过 MQTT smartcar/system/status 上报，'
     '由 ESP32-P4 终端和 Python 上位机统一订阅展示。'
     '状态机设计使小车、摄像头、语音模块不再孤立响应命令，'
     '而是在统一业务流程中按顺序协同工作。')
sp(doc)

h2(doc, '2.1  三大主状态')
tbl(doc,
    ['状态', 'MQTT state 字段', '系统行为', 'Hub 界面'],
    [
        ['巡航状态\ncruise',
         'cruise',
         '小车低速巡航；摄像头、语音保持正常待机',
         '蓝色主题 / 全屋安全 / 待命中'],
        ['异常状态\nabnormal',
         'abnormal_running\nabnormal_ready',
         '小车前往异常地点；摄像头切换图传扫描；语音进行现场问询',
         '红色主题 / 异常复核 / 前往现场'],
        ['返航状态\nreturn_home',
         'return_running',
         '小车按反向路线回起点；关闭摄像头图传与语音联动',
         '蓝色主题 / 全屋安全 / 返航中'],
    ],
    widths=[1.3, 1.5, 2.8, 2.0]
)
sp(doc)

h2(doc, '2.2  卫生间异常完整执行序列（演示主链路）')
body(doc, '以卫生间异常为例，从触发到恢复共 12 步：')
sp(doc)

tbl(doc,
    ['步骤', '执行内容', '通信方式', '状态上报'],
    [
        ['① 触发', 'CSI-WiFi 或上位机检测到卫生间异常，发布 abnormal bathroom 命令',
         'MQTT → smartcar/cmd', '—'],
        ['② 静默语音', 'SmartCar 串口通知 VoiceMode 进入静默（system_silent）',
         'UART', '—'],
        ['③ 静默摄像头', 'SmartCar 通过 MQTT 通知 Camera 静默（silent）',
         'MQTT → cam/jpeg/control', '—'],
        ['④ 前往现场', '小车执行卫生间路线：前进 10s → 左转 700ms → 前进 5s',
         '本地电机控制', 'abnormal_running'],
        ['⑤ 到达', '小车上报 abnormal_ready；串口通知 Voice 开始问询（system_inquiry）',
         'MQTT + UART', 'abnormal_ready'],
        ['⑥ 开始图传', 'Camera 收到 start_scan，切换 MQTT 图传，开始水平扫描（10°~170°）',
         'MQTT → cam/jpeg/control', '—'],
        ['⑦ 人脸跟随', 'Camera 检测到人脸，切换为跟随模式，持续发布 JPEG 帧',
         'MQTT cam/jpeg', '—'],
        ['⑧ 语音问询', '语音模块播放"您是否感到不舒服？"，接入云端语音模型与老人对话',
         'UART + 云端语音', '—'],
        ['⑨ 人工确认', '终端 / 上位机人工判断后发送返航命令',
         'MQTT → smartcar/cmd', '—'],
        ['⑩ 关闭联动', '小车串口通知 Voice: system_off；MQTT 通知 Camera: off',
         'UART + MQTT', '—'],
        ['⑪ 返航', '小车执行返航路线：掉头 → 前进 5s → 右转 → 前进 10s',
         '本地电机控制', 'return_running'],
        ['⑫ 恢复巡航', '小车到达起点，上报 cruise，系统恢复全屋安全状态',
         'MQTT', 'cruise'],
    ],
    widths=[0.55, 3.2, 2.0, 1.5]
)
sp(doc)

h2(doc, '2.3  状态切换与终端界面对应')
body(doc, 'ESP32-P4 终端订阅 smartcar/system/status，根据 state 字段实时更新 UI：')
sp(doc)

tbl(doc,
    ['SmartCar 上报 state', '终端界面变化', '控制按钮变化'],
    [
        ['cruise',           '蓝色主题 / 顶栏"全屋安全" / 状态标题"待命中"',           '"巡航"按钮高亮，"返航"按钮禁用'],
        ['abnormal_running', '红色主题 / 顶栏"异常复核" / 状态标题"前往现场"',         '地点异常按钮可用'],
        ['abnormal_ready',   '红色主题 / 顶栏"异常复核" / 状态标题"就绪确认"',         '"返航"按钮高亮可用'],
        ['return_running',   '蓝色主题 / 顶栏"全屋安全" / 状态标题"返航中"',           '"停止"按钮可用'],
    ],
    widths=[1.8, 3.2, 2.5]
)
sp(doc)

h2(doc, '2.4  各模块在状态机中的角色分工')
tbl(doc,
    ['模块', '巡航阶段', '异常行驶阶段', '到达现场阶段', '返航阶段'],
    [
        ['SmartCar（小车）',  '低速巡航 / 避障',  '执行地点路线',       '等待确认，上报 ready', '执行反向路线'],
        ['Camera（摄像头）',  'LCD 本地显示',     '静默，停止图传',      '扫描跟随，发布 JPEG',  '停止图传'],
        ['Voice（语音）',     '正常待机',         '静默 (system_silent)','问询老人 (inquiry)',    '关闭 (system_off)'],
        ['ESP32-P4 终端',     '显示概况 / 遥控',  '显示红色异常模式',    '显示图传 / 事件日志',  '显示返航状态'],
        ['Python 上位机',     '状态监控 / 调试',  'MJPEG 视频流',        '报警通知',             '状态恢复提示'],
    ],
    widths=[1.3, 1.4, 1.5, 1.7, 1.3]
)
sp(doc)

# ── save ──────────────────────────────────────────────────────────
out = 'f:/test/esp/che/esp_brookesia_phone/演示讲解_终端界面与状态机.docx'
doc.save(out)
print(f'Saved: {out}')
